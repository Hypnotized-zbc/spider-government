#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
重庆市政府采购招标公告爬虫
==========================
目标: https://www.cq.gov.cn/zwgk/zfxxgkml/wlzcxx/zbtb/zcztb/zbgg/
抓取近一个月内的所有招标公告，每份保存为 md + docx，最后打包 zip。

依赖: requests, beautifulsoup4, python-docx
详情页: www.cqggzy.com 有 JS 反爬(521 + cookie 挑战)。
  方案: Windows Edge 无头 + CDP 常驻，逐条导航渲染取 DOM（Edge 只启动一次）。
  前置: 先运行 tools/start_edge_cdp.ps1 启动 Edge。
用法: python3 spider_newest.py
"""
import datetime
import json
import math
import os
import re
import shutil
import subprocess
import sys
import time
import zipfile
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

BASE = "https://www.cq.gov.cn/zwgk/zfxxgkml/wlzcxx/zbtb/zcztb/zbgg"
TOOLS = Path(__file__).parent / "tools"
EDGE_CDP_PS1 = TOOLS / "fetch_via_edge.ps1"
DOWNLOAD_PS1 = TOOLS / "download_v2.ps1"
HEADERS = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
CUTOFF = datetime.date.today() - datetime.timedelta(days=31)  # 近一个月

OUT_DIR = Path(__file__).parent / "output"
DATA_DIR = OUT_DIR / "zbgg_data"
MD_DIR = DATA_DIR / "md"
DOCX_DIR = DATA_DIR / "docx"
ATTACH_DIR = DATA_DIR / "attachments"
META_PATH = DATA_DIR / "meta.json"
FAILED_PATH = DATA_DIR / "failed.json"


def ps1_path() -> str:
    """把 ps1 路径转成 Windows 可执行路径。"""
    return subprocess.run(["wslpath", "-w", str(EDGE_CDP_PS1)],
                          capture_output=True, text=True).stdout.strip()


# ---------- 列表页 ----------

def fetch_list(url: str, tries: int = 3) -> str | None:
    for i in range(tries):
        try:
            r = requests.get(url, headers=HEADERS, timeout=20)
            r.encoding = r.apparent_encoding
            return r.text
        except Exception as e:
            print(f"  列表重试 {i+1}: {e}", flush=True)
            time.sleep(3)
    return None


def parse_list(html: str) -> list[dict]:
    soup = BeautifulSoup(html, "html.parser")
    items = []
    for li in soup.select(".list-warp li"):
        a = li.find("a", class_="ellipsis")
        sp = li.find("span")
        if not a or not sp:
            continue
        try:
            d = datetime.date.fromisoformat(sp.get_text(strip=True))
        except ValueError:
            continue
        items.append({
            "title": a.get_text(strip=True) or a.get("title", ""),
            "url": a["href"],
            "date": sp.get_text(strip=True),
            "date_obj": d,
        })
    return items


def collect_items() -> list[dict]:
    """翻页收集近一个月内条目，遇到整页早于 cutoff 即停。"""
    items = []
    page = 0
    while page < 50:
        url = f"{BASE}/index.html" if page == 0 else f"{BASE}/index_{page}.html"
        print(f"列表页 {page}: {url}", flush=True)
        html = fetch_list(url)
        if not html:
            print(f"  页{page} 获取失败，跳过", flush=True)
            page += 1
            continue
        page_items = parse_list(html)
        if not page_items:
            print(f"  页{page} 无条目，停止", flush=True)
            break
        kept = [it for it in page_items if it["date_obj"] >= CUTOFF]
        items.extend(kept)
        print(f"  本页 {len(kept)}/{len(page_items)} 条在近一月内 (累计 {len(items)})", flush=True)
        if len(kept) < len(page_items):
            print("  已到近一个月边界，停止翻页", flush=True)
            break
        page += 1
        time.sleep(1)
    return items


# ---------- Edge CDP 生命周期 ----------

def edge_alive() -> bool:
    """检查 Edge CDP 是否存活。"""
    try:
        out = subprocess.run(
            ["powershell.exe", "-NoProfile", "-Command",
             "try { (Invoke-RestMethod -Uri 'http://127.0.0.1:9222/json/version' -TimeoutSec 3).Browser } catch { '' }"],
            capture_output=True, timeout=20)
        return bool(out.stdout.decode("utf-8", errors="ignore").strip())
    except Exception:
        return False


def start_edge() -> bool:
    """启动 Edge CDP（Windows 侧）。"""
    ps1 = TOOLS / "start_edge_cdp.ps1"
    if not ps1.exists():
        print("    缺少 start_edge_cdp.ps1", flush=True)
        return False
    win = subprocess.run(["wslpath", "-w", str(ps1)],
                         capture_output=True, text=True).stdout.strip()
    try:
        proc = subprocess.run(
            ["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass",
             "-File", win],
            capture_output=True, timeout=60)
        return edge_alive()
    except Exception as e:
        print(f"    Edge 启动异常: {e}", flush=True)
        return False


def ensure_edge() -> bool:
    """确保 Edge CDP 存活，必要时自动启动。"""
    if edge_alive():
        return True
    print("    Edge CDP 未运行，自动启动…", flush=True)
    return start_edge()


# ---------- 详情页（Edge CDP 渲染） ----------

def fetch_detail(url: str, tries: int = 3) -> str | None:
    """用 Edge CDP 导航渲染详情页；失败自动重启 Edge 并重试。"""
    for attempt in range(1, tries + 1):
        if not edge_alive():
            if not start_edge():
                time.sleep(3)
                continue
        cmd = ["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass",
               "-File", ps1_path(), url]
        try:
            proc = subprocess.run(cmd, capture_output=True, timeout=60)
            html = proc.stdout.decode("utf-8", errors="ignore")
            if html and len(html) >= 2000 and "__jsl_clearance" not in html:
                return html
        except Exception as e:
            print(f"    Edge 渲染异常: {e}", flush=True)
        print(f"    详情获取失败（第 {attempt} 次），重启 Edge 重试…", flush=True)
        time.sleep(3)
    return None


# ---------- 附件下载（Edge CDP，仅 PDF） ----------

ATTACH_BASE = ("https://ggzydl.cqggzy.com/CQTPBidder/jsgcztbmis2/pages/"
               "zbfilelingqu_hy/cQZBFileDownAttachAction.action?cmd=download")


def attach_url(att: dict) -> str:
    return (f"{ATTACH_BASE}&AttachGuid={att['guid']}"
            f"&FileCode={att['code']}&ClientGuid={att['client_guid']}")


def is_pdf(att: dict) -> bool:
    return att.get("name", "").lower().endswith(".pdf")


def download_attachment(att: dict, save_dir: Path, index: int = 0,
                        tries: int = 3) -> Path | None:
    """用 Edge CDP 下载单个附件到 save_dir，返回文件路径；失败重试后返回 None。"""
    save_dir.mkdir(parents=True, exist_ok=True)
    win_tmp = windows_temp_dir()
    win_tmp.mkdir(parents=True, exist_ok=True)
    url = attach_url(att)
    dl_ps1 = subprocess.run(["wslpath", "-w", str(DOWNLOAD_PS1)],
                            capture_output=True, text=True).stdout.strip()
    win_tmp_win = subprocess.run(["wslpath", "-w", str(win_tmp)],
                                 capture_output=True, text=True).stdout.strip()
    cmd = ["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass",
           "-File", dl_ps1, url, win_tmp_win]
    for attempt in range(1, tries + 1):
        # 清空临时目录
        for f in win_tmp.iterdir():
            if f.is_file():
                f.unlink()
        try:
            proc = subprocess.run(cmd, capture_output=True, timeout=90)
            files = [f for f in win_tmp.iterdir() if f.is_file()]
            if files:
                src = files[0]
                # 校验 PDF 头
                if src.read_bytes()[:4] == b"%PDF":
                    safe = safe_name(att["name"]) or f"attach_{index}"
                    dst = save_dir / safe
                    shutil.move(str(src), str(dst))
                    return dst
                print(f"    附件非有效 PDF: {att['name']}", flush=True)
            else:
                print(f"    附件下载无文件（第 {attempt} 次）: {att['name']}", flush=True)
        except Exception as e:
            print(f"    附件下载异常（第 {attempt} 次）: {att['name']}: {e}", flush=True)
        time.sleep(2)
    return None


def download_pdf_attachments(attachments: list[dict], save_dir: Path) -> list[dict]:
    """下载附件中所有 PDF，返回成功清单 [{name, path}]。"""
    pdfs = [a for a in attachments if is_pdf(a)]
    if not pdfs:
        return []
    print(f"    发现 {len(pdfs)} 个 PDF 附件: {[a['name'] for a in pdfs]}", flush=True)
    done = []
    for i, att in enumerate(pdfs):
        path = download_attachment(att, save_dir, index=i)
        if path:
            done.append({"name": att["name"], "path": str(path)})
        time.sleep(0.5)
    return done


def _parse_runs(node, runs: list):
    """递归解析节点为 run 列表，记录下划线。"""
    for child in node.children:
        if isinstance(child, str):
            t = child.strip()
            if t:
                runs.append({"text": t, "underline": False})
        elif child.name == "br":
            runs.append({"text": "\n", "underline": False})
        else:
            style = child.get("style", "") or ""
            underline = child.name == "u" or "underline" in style
            sub = []
            _parse_runs(child, sub)
            for r in sub:
                r["underline"] = r["underline"] or underline
                runs.append(r)


def extract_detail(html: str) -> dict:
    soup = BeautifulSoup(html, "html.parser")
    wrapper = soup.select_one(".detail-wrapper")
    if wrapper:
        h1s = wrapper.find_all("h1")
        title = h1s[0].get_text(strip=True) if h1s else "无标题"
        # 项目编号：h1 形如 "项目编号：xxx"
        code = ""
        for h in h1s:
            t = h.get_text(strip=True)
            if t.startswith("项目编号"):
                code = t.split("：", 1)[-1].strip()
                break
        # 信息时间：wrapper 中文本 "信息时间：yyyy-mm-dd"
        info_time = ""
        for el in wrapper.find_all(string=lambda s: s and "信息时间：" in s):
            m = re.search(r"信息时间：(\S+)", el)
            if m:
                info_time = m.group(1)
                break
    else:
        title_el = soup.select_one("h1") or soup.select_one(".gonggaotitle")
        title = title_el.get_text(strip=True) if title_el else "无标题"
        code = ""
        info_time = ""
    # 正文：按 app-detail 内子块顺序，识别 titleStr(小标题) + content(正文)
    paragraphs = []
    texts = []
    app = soup.select_one(".app-detail")
    blocks = list(app.children) if app else []
    if not blocks:
        blocks = soup.select(".content")
    for el in blocks:
        if not hasattr(el, "name") or not el.name:
            continue
        cls = el.get("class", []) or []
        if "hide" in cls:
            continue  # 页面隐藏块，不输出
        for tag in el.find_all(["script", "style", "button", "meta"]):
            tag.decompose()
        ts = el.find(class_="titleStr")
        if ts:
            runs = []
            _parse_runs(ts, runs)
            if runs:
                paragraphs.append({"type": "heading", "runs": runs})
                texts.append("".join(r["text"] for r in runs))
        cs = el.find(class_="content")
        if cs:
            ps = cs.find_all("p")
            if ps:
                for p in ps:
                    runs = []
                    _parse_runs(p, runs)
                    if runs:
                        paragraphs.append({"type": "body", "runs": runs})
                        texts.append("".join(r["text"] for r in runs))
            else:
                runs = []
                _parse_runs(cs, runs)
                if runs:
                    paragraphs.append({"type": "body", "runs": runs})
                    texts.append("".join(r["text"] for r in runs))
        # 块内表格（联系方式表等位于 titleStr 块内、无 content class）
        for tbl in el.find_all("table"):
            rows = []
            for tr in tbl.find_all("tr"):
                cells = [td.get_text(" ", strip=True) for td in tr.find_all(["td", "th"])]
                cells = [c for c in cells if c]
                if cells:
                    rows.append(cells)
            if rows:
                paragraphs.append({"type": "table", "rows": rows})
                texts.append(" | ".join(" | ".join(r) for r in rows))
        # 独立 table 块（投标保证金账号表等）
        if el.name == "table":
            rows = []
            for tr in el.find_all("tr"):
                cells = [td.get_text(" ", strip=True) for td in tr.find_all(["td", "th"])]
                cells = [c for c in cells if c]
                if cells:
                    rows.append(cells)
            if rows:
                paragraphs.append({"type": "table", "rows": rows})
                texts.append(" | ".join(" | ".join(r) for r in rows))
    # 附件：解析 .ewb-blue-a 的 downloadAttach 链接
    attachments = []
    for a in soup.select(".ewb-blue-a"):
        name = a.get_text(strip=True)
        onclick = a.get("onclick", "") or ""
        m = re.search(r"downloadAttach\('([^']+)','([^']+)','([^']+)'\)", onclick)
        if m and name:
            attachments.append({
                "name": name,
                "guid": m.group(1),
                "code": m.group(2),
                "client_guid": m.group(3),
            })
    return {"title": title, "code": code, "info_time": info_time,
            "paragraphs": paragraphs, "body": "\n\n".join(texts),
            "attachments": attachments}


# ---------- 输出 ----------

def to_markdown(item: dict, detail: dict) -> str:
    """格式：标题行 + 项目编号行 + 信息时间行 + 正文。"""
    lines = [f"项目标题：{detail['title']}"]
    if detail.get("code"):
        lines.append(f"项目编号：{detail['code']}")
    if detail.get("info_time"):
        lines.append(f"信息时间：{detail['info_time']}")
    lines.append("")
    lines.append(detail["body"])
    return "\n".join(lines) + "\n"


def to_docx(md_path: Path, docx_path: Path, detail: dict | None = None):
    """按网页排版生成 docx：标题居中加粗、元信息行、正文宋体/下划线/缩进。"""
    doc = Document()

    def set_cn_font(run, size_pt: float, bold: bool = False):
        run.font.size = Pt(size_pt)
        run.font.name = "宋体"
        run.font.bold = bold
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "宋体")

    if detail is not None:
        # 项目标题：居中、加粗（网页 h1 24px≈18pt）
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cn_font(p.add_run(detail["title"]), 18, bold=True)
        # 项目编号：居中
        if detail.get("code"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cn_font(p.add_run(f"项目编号：{detail['code']}"), 18)
        # 信息时间：居中
        if detail.get("info_time"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cn_font(p.add_run(f"信息时间：{detail['info_time']}"), 14)
        # 正文：按类型输出；heading 小标题加粗不缩进，body 宋体缩进行距下划线，table 转 Word 表格
        for para in detail.get("paragraphs", []):
            ptype = para.get("type")
            if ptype == "table":
                rows = para.get("rows", [])
                if not rows:
                    continue
                ncols = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncols)
                tbl.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(ncols):
                        cell = tbl.cell(ri, ci)
                        cell.text = ""
                        cell.paragraphs[0].add_run(row[ci] if ci < len(row) else "")
                doc.add_paragraph()
                continue
            is_head = ptype == "heading"
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.line_spacing = 1.5
            pf.space_after = Pt(0)
            if not is_head:
                pf.first_line_indent = Pt(24)  # 2 字符 ≈ 24pt
            for r in para.get("runs", []):
                run = p.add_run(r["text"])
                set_cn_font(run, 12, bold=is_head)
                if r.get("underline"):
                    run.underline = True
    else:
        # 回退：纯文本 md
        for line in md_path.read_text(encoding="utf-8").splitlines():
            if line.strip():
                p = doc.add_paragraph()
                set_cn_font(p.add_run(line.strip()), 12)
    doc.save(docx_path)


def make_zip(zip_path: Path, folder: Path):
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in sorted(folder.rglob("*")):
            if f.is_file():
                zf.write(f, f.relative_to(folder))


def make_full_zip(zip_path: Path, docx_dir: Path, att_dir: Path):
    """打包 docx 与 PDF 附件，docx 在根，附件按公告名目录。"""
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in sorted(docx_dir.glob("*.docx")):
            zf.write(f, f.name)
        for d in sorted(att_dir.iterdir()):
            if not d.is_dir():
                continue
            for f in sorted(d.rglob("*")):
                if f.is_file():
                    zf.write(f, f"{d.name}/{f.name}")


def make_report_zip(zip_path: Path, folder_name: str,
                    docx_file: Path | None, att_dir: Path | None):
    """每份报告一个 zip：内部一个以报告标题命名的文件夹，含 docx + PDF 附件。"""
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        if docx_file and docx_file.exists():
            zf.write(docx_file, f"{folder_name}/{docx_file.name}")
        if att_dir and att_dir.is_dir():
            for f in sorted(att_dir.rglob("*")):
                if f.is_file():
                    zf.write(f, f"{folder_name}/{f.name}")


def make_summary_zip(zip_path: Path, records: list[dict]):
    """按本次报告记录打包：每份一个文件夹（docx + PDF 附件）。

    只打包传入的 records（本次爬取范围），不再扫描整个 docx 目录，
    避免历史累积的报告反复出现在压缩包里。
    """
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for rec in records:
            docx_file = Path(rec["docx"])
            if not docx_file.exists():
                continue
            folder_name = docx_file.stem  # 报告标题文件夹名
            zf.write(docx_file, f"{folder_name}/{docx_file.name}")
            for att in rec.get("attachments", []):
                ap = Path(att["path"])
                if ap.exists():
                    zf.write(ap, f"{folder_name}/{ap.name}")


def desktop_path() -> Path:
    """获取 Windows 桌面路径（自动发现，不依赖特定用户名）。"""
    # 1. PowerShell 获取
    try:
        out = subprocess.run(
            ["powershell.exe", "-NoProfile", "-Command",
             "[Environment]::GetFolderPath('Desktop')"],
            capture_output=True, timeout=20)
        p = out.stdout.decode("utf-8", errors="ignore").strip()
        if p and "\\" in p:
            p = p.replace("\\", "/")
            drive = p[0].lower()
            cand = Path(f"/mnt/{drive}{p[2:]}")
            if cand.is_dir():
                return cand
    except Exception:
        pass
    # 2. 扫描 /mnt/c/Users/*/Desktop（自动发现，不写死用户名）
    try:
        for u in Path("/mnt/c/Users").iterdir():
            if u.is_dir():
                cand = u / "Desktop"
                if cand.is_dir():
                    return cand
    except Exception:
        pass
    # 3. 兜底：输出目录下的 desktop 子目录
    d = OUT_DIR / "desktop"
    d.mkdir(parents=True, exist_ok=True)
    return d


def windows_temp_dir() -> Path:
    """获取 Windows 可写的临时目录（WSL 视角），用于附件中转。"""
    try:
        out = subprocess.run(
            ["powershell.exe", "-NoProfile", "-Command",
             "[System.IO.Path]::GetTempPath()"],
            capture_output=True, timeout=20)
        p = out.stdout.decode("utf-8", errors="ignore").strip()
        if p and "\\" in p:
            p = p.replace("\\", "/")
            drive = p[0].lower()
            cand = Path(f"/mnt/{drive}{p[2:]}")
            if cand.is_dir():
                return cand / "zbgg_dl"
    except Exception:
        pass
    # 兜底：项目 output 下
    d = OUT_DIR / "tmp_dl"
    d.mkdir(parents=True, exist_ok=True)
    return d


# ---------- HTML 统计报表 ----------

def parse_md_fields(md_path: Path) -> dict:
    """从 md 文件前几行解析 标题/编号/信息时间。"""
    txt = md_path.read_text(encoding="utf-8")
    title = code = info_time = ""
    for line in txt.splitlines()[:5]:
        if line.startswith("项目标题：") and not title:
            title = line[len("项目标题："):].strip()
        elif line.startswith("项目编号：") and not code:
            code = line[len("项目编号："):].strip()
        elif line.startswith("信息时间：") and not info_time:
            info_time = line[len("信息时间："):].strip()
    return {"title": title, "code": code, "info_time": info_time}


def extract_key_info(body: str) -> dict:
    """从正文提取 投资金额/工期。优先总投资金额，其次合同估算金额。"""
    info = {"investment": "", "duration": ""}
    # 投资金额：总投资金额 / 总投资 / 合同估算金额 / 估算金额（兼容"含税"等前缀）
    for pat in [r"总投资金额[：:]\s*(?:含税|不含税)?\s*([\d.,]+)\s*万元",
                r"总投资[约为]?[：:]\s*(?:含税|不含税)?\s*([\d.,]+)\s*万元",
                r"合同估算金额[：:]\s*(?:含税|不含税)?\s*([\d.,]+)\s*万元",
                r"估算金额[：:]\s*(?:含税|不含税)?\s*([\d.,]+)\s*万元"]:
        m = re.search(pat, body)
        if m:
            info["investment"] = f"{m.group(1).replace(',', '')} 万元"
            break
    # 工期：工期要求 / 服务期限 / 总工期
    for pat in [r"工期要求[：:]?\s*([^。\n；]+)",
                r"总工期[：:]?\s*([^。\n；]+)",
                r"服务期限[：:]?\s*([^。\n；]+)",
                r"工期[约为]?[：:]?\s*([^。\n；]+)"]:
        m = re.search(pat, body)
        if m:
            info["duration"] = m.group(1).strip()
            break
    return info


# ---------- 价值评分（按公司画像） ----------

# 招标类型 → 分值（总承包/施工通常体量大，咨询/监理次之）
TYPE_VALUE = [
    (20, ["EPC", "工程总承包", "设计施工", "设计采购施工"]),
    (16, ["施工总承包", "工程施工", "标段施工", "施工招标"]),
    (12, ["全过程咨询", "项目管理"]),
    (8,  ["监理", "勘察设计", "初步设计", "施工图设计"]),
]
# 业务方向关键词（默认画像用，按自己关心的领域增删，命中越多价值越高）
INTEREST_KEYWORDS = [
    "老旧小区", "学校", "医院", "产业园", "数据中心", "智慧",
    "市政", "道路", "污水", "给排水", "绿化", "装修", "托育",
]
# 资质等级次序：一级 > 二级 > 三级 > 不分等级(0)
QUAL_LEVEL = {"一级": 3, "二级": 2, "三级": 1}

PROFILE_DIR = Path(__file__).parent / "company_profiles"

DEFAULT_PROFILE = {
    "name": "默认（综合视角）",
    "desc": "未配置公司画像时的通用评分",
    "keywords": {k: 1 for k in INTEREST_KEYWORDS},
    "exclude_keywords": [],
    "qualifications": [],
    "budget_range": None,
    "regions": [],
    "weights": {"money": 0.30, "type": 0.15, "attach": 0.15,
                "kw": 0.20, "region": 0.10, "qual": 0.10},
}


def load_profiles() -> list[dict]:
    """读取 company_profiles/*.json，未配置时返回默认画像。"""
    profiles = [dict(DEFAULT_PROFILE)]
    if not PROFILE_DIR.is_dir():
        return profiles
    for f in sorted(PROFILE_DIR.glob("*.json")):
        try:
            p = json.loads(f.read_text(encoding="utf-8"))
        except Exception as e:
            print(f"  画像解析失败 {f.name}: {e}", flush=True)
            continue
        p.setdefault("name", f.stem)
        p.setdefault("desc", "")
        p.setdefault("keywords", {})
        p.setdefault("exclude_keywords", [])
        p.setdefault("qualifications", [])
        p.setdefault("budget_range", None)
        p.setdefault("regions", [])
        p.setdefault("weights", dict(DEFAULT_PROFILE["weights"]))
        profiles.append(p)
    return profiles


def _qual_stem(q: str) -> str:
    """资质主干词：去掉等级与后缀，如 '建筑工程施工总承包二级' → '建筑工程施工总承包'。"""
    return re.sub(r"[一二三四]级|及以上|及以下|不分等级", "", q).strip()


def qual_mismatch(body: str, qualifications: list[str]) -> tuple[str, float]:
    """资质匹配：返回 (说明, 资质分 0-10)。

    - 公司未填资质：中性 6 分
    - 正文要求与公司资质主干词命中且等级不高于公司：10 分
    - 命中但正文要求等级更高：0 分（资质不符）
    - 正文无资质要求：7 分
    - 正文有资质要求但无命中：2 分
    """
    if not qualifications:
        return "未配置资质", 6.0
    if "资质" not in body and "总承包" not in body and "专业承包" not in body:
        return "正文未要求资质", 7.0
    for q in qualifications:
        stem = _qual_stem(q)
        if stem in body:
            m = re.search(re.escape(stem) + r"[^，。；\n]*?([一二三四]级)", body)
            req_lv = QUAL_LEVEL.get(m.group(1), 0) if m else 0
            own_lv = QUAL_LEVEL.get(re.search(r"([一二三四]级)", q).group(1), 0) if re.search(r"([一二三四]级)", q) else 0
            if req_lv > own_lv:
                return f"需{stem}（高于本公司）", 0.0
            return f"资质匹配（{stem}）", 10.0
    return "资质方向不符", 2.0


def region_score(body: str, regions: list[str]) -> tuple[str, float]:
    """地区匹配 0-10：正文建设地点命中关注区域→10，未写地点→5，未命中→0。"""
    m = re.search(r"建设地点[：:]\s*([^。\n；]+)", body)
    loc = m.group(1).strip() if m else ""
    if not regions:
        return "未配置地区", 5.0
    if not loc:
        return "正文未写建设地点", 5.0
    for r in regions:
        if r in loc:
            return f"建设地点在{r}", 10.0
    return f"建设地点不在关注区域（{loc[:12]}）", 0.0


def budget_score(money_v: float, budget_range) -> tuple[str, float]:
    """预算区间匹配：区间内按 log 给分（上限30），超出上限→0，低于下限→低分。"""
    if budget_range is None:
        return "", min(30.0, 7.5 * math.log10(money_v)) if money_v > 0 else 0.0
    lo, hi = budget_range
    if money_v <= 0:
        return "未写明金额", 0.0
    if money_v > hi:
        return f"超出预算上限（{hi}万）", 0.0
    if money_v < lo:
        return f"低于预算下限（{lo}万）", min(10.0, 7.5 * math.log10(money_v))
    return f"在预算区间内", min(30.0, 7.5 * math.log10(money_v))


def value_components(title: str, body: str, att_names: str,
                     profile: dict | None = None) -> dict:
    """按公司画像从标题+正文+附件名计算价值分量，返回 0-100 总分与明细。

    维度（满分=权重×100）：金额30 / 类型15 / 附件15 / 关键词20 / 地区10 / 资质10。
    """
    profile = profile or DEFAULT_PROFILE
    w = profile.get("weights", DEFAULT_PROFILE["weights"])
    full = f"{title} {body}"

    # 金额
    money_v = 0.0
    for pat in [r"总投资金额[：:]\s*(?:含税|不含税)?\s*([\d.,]+)\s*万元",
                r"总投资[约为]?[：:]\s*(?:含税|不含税)?\s*([\d.,]+)\s*万元",
                r"合同估算金额[：:]\s*(?:含税|不含税)?\s*([\d.,]+)\s*万元",
                r"估算金额[：:]\s*(?:含税|不含税)?\s*([\d.,]+)\s*万元"]:
        m = re.search(pat, body)
        if m:
            try:
                money_v = float(m.group(1).replace(",", ""))
            except ValueError:
                pass
            break
    money_note, money_s = budget_score(money_v, profile.get("budget_range"))

    # 类型
    type_s = 4
    for score, kws in TYPE_VALUE:
        if any(kw in full for kw in kws):
            type_s = score
            break
    type_s = min(type_s, 15)

    # 附件
    att_s = 0.0
    if "工程量清单" in att_names or "图纸" in att_names:
        att_s = 15.0
    elif att_names:
        att_s = 8.0

    # 关键词：按画像加权命中
    kws_map = profile.get("keywords", {})
    hit = [kw for kw in kws_map if kw in full]
    kw_s = min(20.0, 2.0 * sum(kws_map.get(kw, 1) for kw in hit))
    # 排除词命中 → 大幅降权
    excl_hit = [kw for kw in profile.get("exclude_keywords", []) if kw in full]
    if excl_hit:
        kw_s = min(kw_s, 5.0) - 5.0  # 命中排除词: 关键词分封顶 0
        if kw_s < 0:
            kw_s = 0.0

    # 地区与资质
    region_note, region_s = region_score(body, profile.get("regions", []))
    qual_note, qual_s = qual_mismatch(body, profile.get("qualifications", []))

    # 加权总分
    total = round(w.get("money", 0.30) * 100 / 30 * money_s
                  + w.get("type", 0.15) * 100 / 15 * type_s
                  + w.get("attach", 0.15) * 100 / 15 * att_s
                  + w.get("kw", 0.20) * 100 / 20 * kw_s
                  + w.get("region", 0.10) * 100 / 10 * region_s
                  + w.get("qual", 0.10) * 100 / 10 * qual_s, 1)

    return {"score": total, "money": money_v, "money_s": money_s,
            "type_s": type_s, "att_s": att_s, "kw_s": kw_s,
            "region_s": region_s, "qual_s": qual_s,
            "kw_hit": hit, "excl_hit": excl_hit,
            "money_note": money_note, "region_note": region_note,
            "qual_note": qual_note, "money_w": round(money_v, 2)}


def score_from_md(md_path: Path, att_names: str = "",
                  profile: dict | None = None) -> dict:
    """从 md 文件 + 附件名计算价值分（供 HTML 报表与 meta 回填用）。"""
    txt = md_path.read_text(encoding="utf-8") if md_path.exists() else ""
    fields = parse_md_fields(md_path)
    return value_components(fields.get("title", ""), txt, att_names, profile)


# ---------- 价值图表（纯 SVG，不依赖外部 CDN） ----------

CHART_COLORS = ["#d35400", "#e67e22", "#f39c12", "#3498db", "#27ae60"]
COMPONENT_NAMES = [("money_s", "金额"), ("type_s", "类型"), ("att_s", "附件"),
                   ("kw_s", "关键词"), ("region_s", "地区"), ("qual_s", "资质")]
COMPONENT_MAX = {"money_s": 30.0, "type_s": 15.0, "att_s": 15.0,
                 "kw_s": 20.0, "region_s": 10.0, "qual_s": 10.0}


def _short_title(title: str, n: int = 14) -> str:
    return title if len(title) <= n else title[:n] + "…"


def _match_notes(r: dict) -> str:
    """拼出匹配原因说明（金额/地区/资质/关键词）。"""
    notes = []
    if r.get("money_note"):
        notes.append(f"金额：{r['money_note']}")
    if r.get("region_note"):
        notes.append(f"地区：{r['region_note']}")
    if r.get("qual_note"):
        notes.append(f"资质：{r['qual_note']}")
    if r.get("kw_hit"):
        notes.append(f"关键词：{'、'.join(r['kw_hit'])}")
    if r.get("excl_hit"):
        notes.append(f"排除词命中：{'、'.join(r['excl_hit'])}")
    return "；".join(notes) if notes else "无特殊说明"


def make_value_charts(rows: list[dict]) -> str:
    """生成价值图表 HTML（总分条形图 + 构成堆叠图 + 匹配明细），无数据返回空串。"""
    if not rows:
        return ""
    srows = sorted(rows, key=lambda r: -r["score"])

    # ---- 图1: 价值总分横向条形图 ----
    H1 = len(srows) * 30 + 46
    parts1 = []
    parts1.append(f'<svg viewBox="0 0 960 {H1}" xmlns="http://www.w3.org/2000/svg" '
                  f'preserveAspectRatio="xMidYMid meet" style="width:100%;height:auto;display:block">')
    parts1.append(f'<text x="8" y="24" font-size="16" font-weight="bold" fill="#1f3a5f">价值总分排行（满分 100）</text>')
    for i, r in enumerate(srows):
        y = 46 + i * 30
        sc = r["score"]
        color = "#c0392b" if sc >= 75 else "#e67e22" if sc >= 60 else "#95a5a6"
        parts1.append(f'<text x="8" y="{y+15}" font-size="12" fill="#333" text-anchor="start">{_short_title(r["title"])}</text>')
        parts1.append(f'<rect x="170" y="{y+2}" width="660" height="20" rx="4" fill="#eef2f7"/>')
        w = max(4, 660 * sc / 100.0)
        parts1.append(f'<rect x="170" y="{y+2}" width="{w:.1f}" height="20" rx="4" fill="{color}"/>')
        parts1.append(f'<text x="838" y="{y+17}" font-size="13" font-weight="bold" fill="#1f3a5f">{sc:.1f}</text>')
    parts1.append("</svg>")

    # ---- 图2: 分量堆叠图（各维度占满分比例） ----
    H2 = len(srows) * 30 + 90
    parts2 = []
    parts2.append(f'<svg viewBox="0 0 960 {H2}" xmlns="http://www.w3.org/2000/svg" '
                  f'preserveAspectRatio="xMidYMid meet" style="width:100%;height:auto;display:block">')
    parts2.append(f'<text x="8" y="24" font-size="16" font-weight="bold" fill="#1f3a5f">价值构成（按维度满分归一）</text>')
    # 图例
    lx = 170
    for (k, name), color in zip(COMPONENT_NAMES, CHART_COLORS):
        parts2.append(f'<rect x="{lx}" y="36" width="12" height="12" fill="{color}"/>')
        parts2.append(f'<text x="{lx+16}" y="46" font-size="12" fill="#333">{name}</text>')
        lx += 66
    for i, r in enumerate(srows):
        y = 62 + i * 30
        parts2.append(f'<text x="8" y="{y+15}" font-size="12" fill="#333">{_short_title(r["title"])}</text>')
        parts2.append(f'<rect x="170" y="{y+2}" width="660" height="20" rx="4" fill="#eef2f7"/>')
        x = 170
        for (k, name), color in zip(COMPONENT_NAMES, CHART_COLORS):
            seg = 660 * r[k] / COMPONENT_MAX[k]
            if seg < 0.5:
                continue
            parts2.append(f'<rect x="{x:.1f}" y="{y+2}" width="{seg:.1f}" height="20" fill="{color}"/>')
            x += seg
        parts2.append(f'<text x="838" y="{y+17}" font-size="13" font-weight="bold" fill="#1f3a5f">{r["score"]:.1f}</text>')
    parts2.append("</svg>")

    # ---- 匹配明细 ----
    notes = []
    for i, r in enumerate(srows, 1):
        notes.append(f"<li><b>{i}. {r['title']}</b>（{r['score']:.1f} 分）<br>"
                     f"<span class=\"note\">{_match_notes(r)}</span></li>")

    return (f'<div class="charts">'
            f'<div class="chart-card"><div class="chart-title">价值高低总览</div>{chr(10).join(parts1)}</div>'
            f'<div class="chart-card"><div class="chart-title">价值构成分析</div>{chr(10).join(parts2)}'
            f'<ul class="notes">{chr(10).join(notes)}</ul></div>'
            f'</div>')


def _profile_desc(p: dict) -> str:
    """公司画像说明（用于报表展示）。"""
    parts = [p.get("desc", "")] if p.get("desc") else []
    kws = p.get("keywords", {})
    if kws:
        parts.append("关键词: " + "、".join(f"{k}({v})" for k, v in kws.items()))
    if p.get("exclude_keywords"):
        parts.append("排除: " + "、".join(p["exclude_keywords"]))
    if p.get("qualifications"):
        parts.append("资质: " + "、".join(p["qualifications"]))
    br = p.get("budget_range")
    if br:
        parts.append(f"预算区间: {br[0]}~{br[1]}万元")
    if p.get("regions"):
        parts.append("关注区域: " + "、".join(p["regions"]))
    return " ｜ ".join(parts) if parts else "无特殊配置"


# ---------- LLM 动态价值报告（DeepSeek） ----------

LLM_BASE_URL = "https://api.deepseek.com/v1"
LLM_MODEL = "deepseek-v4-flash"
LLM_TIMEOUT = 120


def load_llm_key() -> str:
    """读取 DeepSeek API key：优先环境变量 DEEPSEEK_API_KEY，其次本地 llm_key.txt。"""
    env = os.environ.get("DEEPSEEK_API_KEY", "").strip()
    if env:
        return env
    key_file = Path(__file__).parent / "llm_key.txt"
    if key_file.exists():
        return key_file.read_text(encoding="utf-8").strip()
    return ""


def llm_chat(system: str, user: str, tries: int = 2) -> str | None:
    """调用 DeepSeek chat/completions，返回助手文本；失败返回 None。"""
    key = load_llm_key()
    if not key:
        print("  未配置 DeepSeek API key（llm_key.txt 或环境变量 DEEPSEEK_API_KEY），跳过 AI 报告", flush=True)
        return None
    payload = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.3,
        "max_tokens": 4000,
        "response_format": {"type": "json_object"},
        "stream": False,
    }
    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
    for i in range(1, tries + 1):
        try:
            r = requests.post(f"{LLM_BASE_URL}/chat/completions",
                              json=payload, headers=headers, timeout=LLM_TIMEOUT)
            if r.status_code == 401:
                print("  DeepSeek API key 无效（401）", flush=True)
                return None
            r.raise_for_status()
            data = r.json()
            return data["choices"][0]["message"]["content"].strip()
        except Exception as e:
            print(f"  LLM 调用失败（第 {i} 次）: {e}", flush=True)
            time.sleep(2)
    return None


def llm_value_report(meta: list[dict], profile: dict) -> dict | None:
    """调用 DeepSeek 动态生成某公司视角的价值报告，返回解析后的 dict 或 None。

    输入：每条公告的 标题/编号/金额/工期/附件名/正文前 300 字；
    输出：{"summary": 总体判断, "items": [{index, score, verdict, reason}], "notes": 建议}
    """
    rows = []
    for i, m in enumerate(meta, 1):
        md_path = Path(m["md"])
        body = md_path.read_text(encoding="utf-8") if md_path.exists() else ""
        # 去掉 md 头部的 项目标题/项目编号/信息时间 行
        body_lines = [l for l in body.splitlines()
                      if not l.startswith(("项目标题：", "项目编号：", "信息时间："))]
        body_clean = re.sub(r"\s+", " ", "\n".join(body_lines)).strip()
        att_names = "、".join(a["name"] for a in m.get("attachments", []))
        rows.append({
            "index": i,
            "title": m.get("title", ""),
            "investment": extract_key_info(body).get("investment", ""),
            "duration": extract_key_info(body).get("duration", ""),
            "attachments": att_names,
            "body": body_clean[:300],
        })
    if not rows:
        return None
    profile_desc = _profile_desc(profile)
    system = ("你是资深工程招投标分析师。根据给定公司的业务画像，"
              "从一批招标公告中判断哪些值得该公司跟进，并给出价值评分与理由。"
              "只输出 JSON，不要多余文字。")
    user = (f"公司画像：{profile_desc}\n\n"
            f"招标公告列表（JSON）：\n{json.dumps(rows, ensure_ascii=False)}\n\n"
            "请输出 JSON，格式严格如下：\n"
            "{\"summary\": \"对这批公告的整体价值判断，100字以内\",\n"
            " \"items\": [{\"index\": 1, \"score\": 0-100整数, "
            "\"verdict\": \"重点跟进|值得关注|一般|忽略\", "
            "\"reason\": \"一句话理由，60字以内\"}, ...],\n"
            " \"notes\": \"给该公司的整体建议，100字以内\"}\n"
            "items 必须覆盖列表中的每一条公告。")
    text = llm_chat(system, user)
    if not text:
        return None
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        # 兼容模型输出被 ```json 包裹的情况
        m = re.search(r"```(?:json)?\s*(\{.*\})\s*```", text, re.S)
        if m:
            try:
                return json.loads(m.group(1))
            except json.JSONDecodeError:
                return None
        return None


def render_llm_report_html(report: dict) -> str:
    """把 LLM 价值报告 dict 渲染成 HTML 片段；解析失败返回空串。"""
    if not report:
        return ""
    summary = report.get("summary", "")
    notes = report.get("notes", "")
    items = report.get("items", [])
    if not isinstance(items, list) or not items:
        return ""
    # 按分数降序
    try:
        items = sorted(items, key=lambda it: -float(it.get("score", 0)))
    except (TypeError, ValueError):
        pass
    verdict_color = {"重点跟进": "#c0392b", "值得关注": "#e67e22",
                     "一般": "#7f8c8d", "忽略": "#bdc3c7"}
    lis = []
    for it in items:
        idx = it.get("index", "-")
        score = it.get("score", "-")
        verdict = it.get("verdict", "")
        reason = it.get("reason", "")
        color = verdict_color.get(verdict, "#7f8c8d")
        lis.append(
            f'<li><b>#{idx}</b> <span style="color:{color}">[{verdict}]</span> '
            f'<b>{score}</b>分 — {reason}</li>')
    parts = ['<div class="llm-card">']
    parts.append('<div class="llm-title">AI 动态价值报告（DeepSeek）</div>')
    if summary:
        parts.append(f'<div class="llm-summary">总体判断：{summary}</div>')
    parts.append('<ol class="llm-items">' + "".join(lis) + "</ol>")
    if notes:
        parts.append(f'<div class="llm-notes">建议：{notes}</div>')
    parts.append("</div>")
    return "\n".join(parts)


def generate_html_report(meta: list[dict], html_path: Path,
                         profiles: list[dict] | None = None) -> Path:
    """根据报告记录生成 HTML 统计报表（单文件多公司下拉切换）。

    每个公司画像生成独立视图块（画像说明 + 价值图表 + 原表按该公司价值降序），
    顶部下拉切换，前端 JS 仅做显示切换，无外部依赖。未配置画像时只用默认视角。
    """
    profiles = profiles if profiles is not None else load_profiles()
    base_rows = []
    for i, m in enumerate(meta, 1):
        md_path = Path(m["md"])
        fields = parse_md_fields(md_path) if md_path.exists() else {}
        body = md_path.read_text(encoding="utf-8") if md_path.exists() else ""
        key = extract_key_info(body)
        att_names = "、".join(a["name"] for a in m.get("attachments", []))
        base_rows.append({
            "index": i,
            "title": fields.get("title") or m["title"],
            "url": m["url"],
            "code": fields.get("code") or "-",
            "info_time": fields.get("info_time") or m["date"],
            "investment": key["investment"] or "-",
            "duration": key["duration"] or "-",
            "att": att_names or "-",
        })

    # 默认视角永远在首位；自定义公司在后，保持 company_profiles 内的文件顺序
    views = []
    options = []
    for pi, p in enumerate(profiles):
        rows = []
        for i, b in enumerate(base_rows, 1):
            md_path = Path(meta[i - 1]["md"])
            att_names = "、".join(a["name"] for a in meta[i - 1].get("attachments", []))
            v = score_from_md(md_path, att_names, p)
            r = dict(b)
            r.update({"score": v["score"], "money_s": v["money_s"], "type_s": v["type_s"],
                      "att_s": v["att_s"], "kw_s": v["kw_s"], "region_s": v["region_s"],
                      "qual_s": v["qual_s"], "kw_hit": v["kw_hit"], "excl_hit": v["excl_hit"],
                      "money_note": v["money_note"], "region_note": v["region_note"],
                      "qual_note": v["qual_note"]})
            rows.append(r)

        # 大模型动态价值报告（失败则回退到规则评分，不影响报表）
        print(f"  生成 AI 价值报告（{p.get('name', '公司')}）…", flush=True)
        llm_report = llm_value_report(meta, p)
        llm_html = render_llm_report_html(llm_report) if llm_report else ""

        srows = sorted(rows, key=lambda r: -r["score"])
        body_rows = []
        for r in srows:
            body_rows.append(f"""<tr>
<td>{r['index']}</td>
<td><a href="{r['url']}" target="_blank">{r['title']}</a></td>
<td>{r['code']}</td>
<td>{r['info_time']}</td>
<td>{r['investment']}</td>
<td>{r['duration']}</td>
<td>{r['att']}</td>
</tr>""")
        views.append(f"""<div class="view" id="view-{pi}"{' style="display:none"' if pi else ''}>
<div class="profile-desc">{_profile_desc(p)}</div>
{llm_html}
{make_value_charts(rows)}
<table>
<thead>
<tr><th>序号</th><th>名称</th><th>项目编号</th><th>信息时间</th><th>投资金额</th><th>工期</th><th>附件</th></tr>
</thead>
<tbody>
{''.join(body_rows)}
</tbody>
</table>
</div>""")
        options.append(f'<option value="{pi}"{" selected" if pi == 0 else ""}>{p.get("name", f"公司{pi}")}</option>')

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>招标报告统计报表（多公司价值视图）</title>
<style>
body {{ font-family: "Microsoft YaHei", "宋体", sans-serif; margin: 30px; background: #f5f7fa; }}
h1 {{ text-align: center; color: #1f3a5f; }}
.meta {{ text-align: center; color: #666; margin-bottom: 20px; }}
.selector {{ text-align: center; margin: 16px 0 8px; }}
.selector select {{ font-size: 16px; padding: 8px 16px; border: 1px solid #d0d7e2; border-radius: 6px; background: #fff; }}
.profile-desc {{ background: #eef2f7; border-radius: 8px; padding: 12px 16px; margin: 8px 0 16px; color: #444; font-size: 13px; }}
table {{ border-collapse: collapse; width: 100%; background: #fff; box-shadow: 0 2px 8px rgba(0,0,0,.08); }}
th, td {{ border: 1px solid #d0d7e2; padding: 10px 12px; font-size: 14px; text-align: left; }}
th {{ background: #1f3a5f; color: #fff; }}
tr:nth-child(even) {{ background: #f8fafc; }}
a {{ color: #1f6feb; text-decoration: none; }}
a:hover {{ text-decoration: underline; }}
.charts {{ display: flex; flex-wrap: wrap; gap: 20px; margin-bottom: 24px; }}
.chart-card {{ flex: 1 1 480px; background: #fff; border-radius: 8px; box-shadow: 0 2px 8px rgba(0,0,0,.08); padding: 16px; }}
.chart-title {{ font-size: 15px; font-weight: bold; color: #1f3a5f; margin-bottom: 8px; }}
.notes {{ margin: 12px 0 0; padding-left: 18px; font-size: 12px; color: #555; }}
.notes li {{ margin-bottom: 6px; }}
.note {{ color: #888; }}
.llm-card {{ background: #fff8f0; border: 1px solid #f0d9b0; border-radius: 8px; padding: 14px 16px; margin: 8px 0 16px; }}
.llm-title {{ font-size: 15px; font-weight: bold; color: #8a4b00; margin-bottom: 8px; }}
.llm-summary {{ font-size: 13px; color: #6b4a1e; margin-bottom: 8px; line-height: 1.6; }}
.llm-items {{ margin: 0 0 8px 20px; padding: 0; font-size: 13px; color: #444; }}
.llm-items li {{ margin-bottom: 5px; line-height: 1.5; }}
.llm-notes {{ font-size: 13px; color: #8a4b00; background: #fdf1e0; border-radius: 6px; padding: 8px 12px; }}
</style>
</head>
<body>
<h1>招标报告统计报表（按公司视角）</h1>
<p class="meta">共 {len(base_rows)} 份报告 · 生成时间 {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} · 下拉切换不同公司的价值判断</p>
<div class="selector">
<label for="company">选择公司视角：</label>
<select id="company" onchange="switchView(this.value)">
{''.join(options)}
</select>
</div>
{''.join(views)}
<script>
function switchView(idx) {{
  document.querySelectorAll('.view').forEach(function(v) {{ v.style.display = 'none'; }});
  var el = document.getElementById('view-' + idx);
  if (el) el.style.display = '';
  try {{ localStorage.setItem('zbgg_view', idx); }} catch (e) {{}}
}}
(function() {{
  var sel = document.getElementById('company');
  var saved = null;
  try {{ saved = localStorage.getItem('zbgg_view'); }} catch (e) {{}}
  if (saved !== null && document.getElementById('view-' + saved)) {{
    sel.value = saved;
    switchView(saved);
  }}
}})();
</script>
</body>
</html>"""
    html_path.write_text(html, encoding="utf-8")
    return html_path


def open_html(html_path: Path):
    """用默认浏览器打开 HTML。"""
    try:
        win = subprocess.run(["wslpath", "-w", str(html_path)],
                             capture_output=True, text=True).stdout.strip()
        subprocess.run(["powershell.exe", "-NoProfile", "-Command",
                        f"Start-Process '{win}'"],
                       capture_output=True, timeout=30)
    except Exception as e:
        print(f"打开 HTML 失败: {e}", flush=True)


# ---------- 主流程 ----------

def safe_name(title: str) -> str:
    return "".join(c if c not in '\\/:*?"<>|' else "_" for c in title)[:80]


def ask_scope(items: list[dict]) -> list[dict]:
    """交互询问爬取范围：日期范围 / 序号范围 / 全部。"""
    print("\n请选择爬取范围：", flush=True)
    print("  1. 按日期范围（例如 2026-07-01 至 2026-08-07）", flush=True)
    print("  2. 按序号范围（例如 第 1 份 至 第 10 份）", flush=True)
    print("  3. 全部（近一个月）", flush=True)
    choice = input("请输入选项 [1/2/3，回车默认 3]：").strip() or "3"

    if choice == "1":
        start = input("开始日期 (YYYY-MM-DD)：").strip()
        end = input("结束日期 (YYYY-MM-DD)：").strip()
        try:
            d1 = datetime.date.fromisoformat(start)
            d2 = datetime.date.fromisoformat(end)
        except ValueError:
            print("日期格式错误，改为爬取全部。", flush=True)
            return items
        if d1 > d2:
            d1, d2 = d2, d1
        picked = [it for it in items if d1 <= it["date_obj"] <= d2]
        print(f"按日期范围 {d1} ~ {d2}，命中 {len(picked)} 份。", flush=True)
        return picked

    if choice == "2":
        start = input("起始序号（第几份）：").strip()
        end = input("结束序号（第几份）：").strip()
        try:
            s = int(start)
            e = int(end)
        except ValueError:
            print("序号格式错误，改为爬取全部。", flush=True)
            return items
        if s < 1 or e > len(items) or s > e:
            print(f"序号超出范围（共 {len(items)} 份），改为爬取全部。", flush=True)
            return items
        picked = items[s - 1:e]
        print(f"按序号 {s} ~ {e}，命中 {len(picked)} 份。", flush=True)
        return picked

    return items


def main(limit: int | None = None):
    print(f"=== 重庆招标公告爬虫 ===\n近一个月: 自 {CUTOFF} 起", flush=True)
    for d in (DATA_DIR, MD_DIR, DOCX_DIR, ATTACH_DIR):
        d.mkdir(parents=True, exist_ok=True)

    # 确保 Edge CDP 存活
    if not ensure_edge():
        print("Edge CDP 无法启动，终止", flush=True)
        return

    # 已有进度（断点续爬）
    done = set()
    if META_PATH.exists():
        for m in json.loads(META_PATH.read_text(encoding="utf-8")):
            done.add(m["url"])
    failed = []
    if FAILED_PATH.exists():
        failed = json.loads(FAILED_PATH.read_text(encoding="utf-8"))

    items = collect_items()
    seen = set()
    items_dedup = []
    for it in items:
        key = (it["title"], it["date"])
        if key not in seen:
            seen.add(key)
            items_dedup.append(it)
    items = items_dedup
    if limit:
        items = items[:limit]
        print(f"（--limit {limit}，仅处理前 {len(items)} 条）", flush=True)
    else:
        items = ask_scope(items)
    print(f"去重后 {len(items)} 条（已完成 {len(done)}，待抓 {len(items)-len(done)}）", flush=True)

    meta = []
    if META_PATH.exists():
        meta = json.loads(META_PATH.read_text(encoding="utf-8"))

    # 本次范围的结果集：历史已爬的直接复用，新爬的实时加入
    cur_meta = []
    for idx, it in enumerate(items, 1):
        if it["url"] in done:
            for m in meta:
                if m["url"] == it["url"]:
                    cur_meta.append(m)
                    break
            continue
        print(f"[{idx}/{len(items)}] {it['date']} {it['title'][:40]}", flush=True)
        html = fetch_detail(it["url"])
        if not html:
            print("    跳过（详情获取失败）", flush=True)
            failed.append(it)
            continue
        detail = extract_detail(html)
        if not detail["body"]:
            print("    跳过（正文为空）", flush=True)
            failed.append(it)
            continue
        name = f"{it['date']}_{safe_name(detail['title'])}"
        md_path = MD_DIR / f"{name}.md"
        docx_path = DOCX_DIR / f"{name}.docx"
        md_path.write_text(to_markdown(it, detail), encoding="utf-8")
        to_docx(md_path, docx_path, detail)
        # 下载 PDF 附件（每份公告一个子目录）
        att_dir = ATTACH_DIR / name
        att_dir.mkdir(parents=True, exist_ok=True)
        pdfs = download_pdf_attachments(detail.get("attachments", []), att_dir)
        att_names = "、".join(a["name"] for a in pdfs)
        v = value_components(detail["title"], detail["body"], att_names)
        rec = {"date": it["date"], "title": detail["title"],
               "url": it["url"], "md": str(md_path), "docx": str(docx_path),
               "attachments": pdfs, "score": v["score"]}
        meta.append(rec)
        cur_meta.append(rec)
        done.add(it["url"])
        # 每 10 条落盘一次进度
        if len(meta) % 10 == 0:
            META_PATH.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
            FAILED_PATH.write_text(json.dumps(failed, ensure_ascii=False, indent=2), encoding="utf-8")
        time.sleep(0.5)

    META_PATH.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    FAILED_PATH.write_text(json.dumps(failed, ensure_ascii=False, indent=2), encoding="utf-8")

    # 汇总打包：只打包本次范围（docx + PDF 附件），不再混入历史报告
    ZIP_DIR = OUT_DIR / "zbgg_zips"
    ZIP_DIR.mkdir(parents=True, exist_ok=True)
    summary_zip = ZIP_DIR / "招标报告汇总.zip"
    desktop_dest = None
    if cur_meta:
        make_summary_zip(summary_zip, cur_meta)
        # 复制到桌面
        try:
            dest = desktop_path() / "招标报告汇总.zip"
            shutil.copy2(str(summary_zip), str(dest))
            desktop_dest = dest
        except Exception as e:
            print(f"复制到桌面失败: {e}", flush=True)
            desktop_dest = None

    print(f"=== 完成: 本次 {len(cur_meta)} 份公告（历史累计 {len(meta)} 份）===", flush=True)
    print(f"md 目录: {MD_DIR}", flush=True)
    print(f"docx 目录: {DOCX_DIR}", flush=True)
    print(f"附件目录: {ATTACH_DIR}", flush=True)
    print(f"汇总压缩包: {summary_zip}", flush=True)
    if desktop_dest:
        print(f"已复制到桌面: {desktop_dest}", flush=True)

    # 生成 HTML 统计报表并打开（只含本次范围）
    if cur_meta:
        html_path = OUT_DIR / "zbgg_report.html"
        generate_html_report(cur_meta, html_path)
        print(f"统计报表: {html_path}", flush=True)
        open_html(html_path)

    if failed:
        print(f"失败 {len(failed)} 条，见 {FAILED_PATH}", flush=True)


if __name__ == "__main__":
    limit = None
    if len(sys.argv) > 1 and sys.argv[1] == "--limit":
        limit = int(sys.argv[2])
    main(limit=limit)
